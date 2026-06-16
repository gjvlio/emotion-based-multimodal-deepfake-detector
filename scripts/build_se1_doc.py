"""
build_se1_doc.py — Generate Software Engineering 1 documentation (Chapter I) for
the DeepSentinel system, faithful to the 2nd Sem AY 2025-2026 format guidelines.

Format spec (from guidelines):
  Paper 8.5 x 11 | Times New Roman 11 | 1.5 line spacing | 1" margins all sides
  Page number bottom center | Fully justified body | Title page elements.

NOTE: bracketed [PLACEHOLDER] tokens are facts the author must supply
(member names, leader, program, dates). Everything else is written from the
actual project.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"D:\Documents\Programming\Thesis_G10\docs\SE1_Documentation_DeepSentinel.docx"

FONT = "Times New Roman"
SIZE = 11

doc = Document()

# ── Global format: Normal style ───────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(SIZE)
# ensure east-asian/complex also TNR
rpr = normal.element.get_or_add_rPr()
rfonts = rpr.get_or_add_rFonts()
rfonts.set(qn("w:ascii"), FONT)
rfonts.set(qn("w:hAnsi"), FONT)
rfonts.set(qn("w:cs"), FONT)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_after = Pt(0)
pf.space_before = Pt(0)

# ── Page setup: 8.5 x 11, 1" margins ──────────────────────────────────────────
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


# ── Helpers ───────────────────────────────────────────────────────────────────
def _set_run(run, bold=False, italic=False, size=SIZE):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    r = run._element
    rpr_ = r.get_or_add_rPr()
    rf = rpr_.get_or_add_rFonts()
    rf.set(qn("w:ascii"), FONT)
    rf.set(qn("w:hAnsi"), FONT)
    rf.set(qn("w:cs"), FONT)


def para(text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False,
         size=SIZE, space_after=6, space_before=0, indent_left=None,
         line_spacing=True):
    p = doc.add_paragraph()
    p.alignment = align
    pf_ = p.paragraph_format
    if line_spacing:
        pf_.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf_.space_after = Pt(space_after)
    pf_.space_before = Pt(space_before)
    if indent_left is not None:
        pf_.left_indent = Inches(indent_left)
    if text:
        run = p.add_run(text)
        _set_run(run, bold=bold, italic=italic, size=size)
    return p


def runs_para(parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6,
              indent_left=None):
    """parts = list of (text, bold, italic) tuples on one paragraph."""
    p = doc.add_paragraph()
    p.alignment = align
    pf_ = p.paragraph_format
    pf_.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf_.space_after = Pt(space_after)
    if indent_left is not None:
        pf_.left_indent = Inches(indent_left)
    for text, bold, italic in parts:
        r = p.add_run(text)
        _set_run(r, bold=bold, italic=italic)
    return p


def heading(text):
    return para(text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True,
                space_before=12, space_after=8)


def _shade_cell(cell, hexcolor):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcpr.append(shd)


def _add_page_number_footer():
    """Bottom-center page number in every section footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # PAGE field
        run = p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)
        _set_run(run)


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
def tcenter(text, bold=False, italic=False, size=SIZE, space_after=0, space_before=0):
    return para(text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=bold, italic=italic,
                size=size, space_after=space_after, space_before=space_before)

tcenter("", space_after=24)
tcenter("POLYTECHNIC UNIVERSITY OF THE PHILIPPINES", bold=True, size=13, space_after=2)
tcenter("College of Computer and Information Sciences", bold=False, size=12, space_after=2)
tcenter("Sta. Mesa, Manila", bold=False, size=12, space_after=48)

tcenter("DeepSentinel: A Multimodal Emotion-Aware", bold=True, size=16, space_after=2)
tcenter("Deepfake Detection System", bold=True, size=16, space_after=60)

tcenter("A Project Presented to the Faculty of the", size=12, space_after=2)
tcenter("College of Computer and Information Sciences", size=12, space_after=2)
tcenter("Polytechnic University of the Philippines", size=12, space_after=2)
tcenter("In Partial Fulfillment of the Requirements in Software Engineering 1", size=12, space_after=48)

tcenter("Submitted by:", bold=True, size=12, space_after=8)
tcenter("[LEADER FULL NAME] — Leader", size=12, space_after=2)
tcenter("[MEMBER 2 FULL NAME]", size=12, space_after=2)
tcenter("[MEMBER 3 FULL NAME]", size=12, space_after=2)
tcenter("[MEMBER 4 FULL NAME]", size=12, space_after=2)
tcenter("[MEMBER 5 FULL NAME]", size=12, space_after=24)

tcenter("Program: [e.g., Bachelor of Science in Computer Science]", size=12, space_after=36)

tcenter("Submitted to:", bold=True, size=12, space_after=8)
tcenter("Assoc. Prof. Melvin C. Roxas", size=12, space_after=2)
tcenter("Software Engineering 1", size=12, space_after=2)
tcenter("2nd Semester, A.Y. 2025–2026", size=12, space_after=2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER I — PROJECT DEFINITION
# ══════════════════════════════════════════════════════════════════════════════
para("CHAPTER I", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=2, space_before=0)
para("PROJECT DEFINITION", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=12)

# ── 1. Introduction ───────────────────────────────────────────────────────────
heading("1. Introduction")
para(
    "DeepSentinel is a multimodal, emotion-aware deepfake detection system designed to "
    "distinguish authentic audio-visual recordings from synthetically manipulated ones. "
    "The proliferation of generative artificial intelligence has made it inexpensive to "
    "fabricate convincing videos in which a person appears to say or feel something they "
    "never did. Conventional detectors search for low-level visual artifacts, which "
    "state-of-the-art generators increasingly suppress. DeepSentinel takes a different "
    "stance grounded in human communication: in a genuine recording, the emotion carried "
    "by a person's voice and the emotion expressed on their face are consistent, whereas "
    "a deepfake—often assembled from mismatched audio and video sources—exhibits "
    "cross-modal emotional incongruence that survives even high-quality synthesis."
)
para(
    "By nature, DeepSentinel is a research-oriented machine-learning system composed of "
    "two cooperating subsystems. The first is a data-generation subsystem that produces a "
    "large, labeled corpus of deepfakes across four escalating manipulation tracks—audio "
    "tampering, lip-sync correction, full-face synthesis, and emotion-mismatch lip "
    "synchronization. The second is a detection subsystem that extracts emotional "
    "representations from each modality and reasons about their agreement."
)
para(
    "Functionally, the system accepts a short video clip and processes it through a "
    "feature-extraction pipeline. A Wav2Vec2 encoder captures acoustic emotion from the "
    "raw waveform while a BERT encoder captures linguistic emotion from the transcribed "
    "speech; together they form a 1,536-dimensional audio-text embedding. A Vision "
    "Transformer (ViT) summarizes the facial emotion across eight expressive keyframes "
    "into a 768-dimensional visual embedding. Two emotion heads independently predict the "
    "emotion of each modality, and their difference yields a six-dimensional mismatch "
    "vector (Delta). In parallel, Compact Bilinear Pooling fuses the two embeddings to "
    "capture fine-grained cross-modal interactions, and a sarcasm head flags linguistic "
    "incongruity. A multilayer-perceptron classifier consumes the fused representation and "
    "outputs the probability that the clip is fake. The system thereby not only renders a "
    "real-or-fake verdict but also exposes the emotional evidence behind that verdict."
)

# ── 2. Project Case ───────────────────────────────────────────────────────────
heading("2. Project Case")

# 2.1 Objectives
runs_para([("2.1. Objectives", True, False)], space_after=6)
runs_para([("General Objective. ", True, False),
           ("To design and develop DeepSentinel, a multimodal emotion-aware deepfake "
            "detection system that identifies manipulated audio-visual media by detecting "
            "emotional incongruence between a speaker's voice and face.", False, False)],
          indent_left=0.3)
runs_para([("Specific Objectives. ", True, False),
           ("Specifically, the project aims:", False, False)], indent_left=0.3)
objectives = [
    "To build a labeled deepfake dataset through a four-track generation pipeline "
    "(audio tampering, lip-sync correction, full-face synthesis, and emotion-mismatch "
    "lip synchronization) that exposes the detector to progressively harder fakes;",
    "To extract and fuse audio-text and visual emotional features using Wav2Vec2, BERT, "
    "and ViT encoders combined through Compact Bilinear Pooling;",
    "To train a multi-task classifier that jointly learns fake-versus-real detection, "
    "per-modality emotion recognition, and sarcasm detection in order to strengthen the "
    "primary detection signal;",
    "To evaluate the model on an unseen benchmark (FakeAVCeleb) and to establish the "
    "statistical significance of its performance against a logistic-regression baseline "
    "using DeLong's test; and",
    "To achieve cross-dataset generalization, targeting an Area Under the ROC Curve "
    "(AUC) of at least 0.70 on manipulation methods never seen during training.",
]
for i, o in enumerate(objectives, 1):
    runs_para([(f"{i}. ", False, False), (o, False, False)], indent_left=0.6)

# 2.2 Scope and Limitation
runs_para([("2.2. Scope and Limitation", True, False)], space_after=6)
para(
    "DeepSentinel covers the full life cycle of deepfake detection—from synthetic data "
    "generation, through feature extraction, to model training and evaluation. The system "
    "is organized into the following modules and sub-modules:",
    indent_left=0.3,
)

modules = [
    ("A. Data Generation Module.",
     "Produces the labeled fake corpus used for training. It comprises four sub-modules: "
     "(1) Track 1 – Audio Tampering, which swaps a clip's audio for emotionally "
     "mismatched speech synthesized with StyleTTS2 and re-voiced with Retrieval-based "
     "Voice Conversion (RVC); (2) Track 2 – Lip-Sync Correction, which additionally "
     "applies Wav2Lip so the lips match the tampered audio; (3) Track 3 – Full-Face "
     "Synthesis, which regenerates the entire talking head with SadTalker; and "
     "(4) Track 4 – Emotion-Mismatch Lip Sync, which uses MuseTalk on MELD dialogue to "
     "pair a face of one emotion with a voice of another."),
    ("B. Preprocessing and Feature Extraction Module.",
     "Converts raw clips into cached feature vectors. Its audio-text sub-module runs "
     "Wav2Vec2 on the waveform and BERT on the speech transcript to form the Z_at "
     "embedding; its visual sub-module performs face detection, keyframe scoring, and ViT "
     "encoding to form the Z_v embedding."),
    ("C. Detection Module.",
     "Performs the real-or-fake decision. Sub-modules include Emotion Head A (audio-text) "
     "and Emotion Head B (visual), the Delta mismatch computation, Compact Bilinear "
     "Pooling, the Sarcasm Head, and the multilayer-perceptron classifier."),
    ("D. Training Module.",
     "Implements two-phase optimization: Phase 1 trains the detection components on "
     "cached features with frozen backbones, and Phase 2 fine-tunes the encoders "
     "end-to-end. Both phases are driven by a multi-task loss that combines detection, "
     "emotion, and sarcasm objectives."),
    ("E. Evaluation Module.",
     "Measures performance through internal accuracy, F1, and AUC; the FakeAVCeleb "
     "cross-dataset benchmark; a logistic-regression baseline; and DeLong's test for "
     "statistical significance."),
]
for title, body in modules:
    runs_para([(title + " ", True, False), (body, False, False)], indent_left=0.6)

para("The project observes the following limitations:", indent_left=0.3, space_before=6)
limitations = [
    "Detection targets short clips of roughly two to eight seconds containing a single "
    "dominant speaker and emotion; multi-speaker or long-form footage is out of scope.",
    "The detector is trained against specific generation methods; entirely novel "
    "manipulation techniques may evade it.",
    "A visible face and audible speech are required—face-occluded or silent clips cannot "
    "be evaluated.",
    "FakeAVCeleb is used strictly as an unseen test set and never for training, in order "
    "to preserve the cross-dataset generalization claim.",
    "Development hardware is limited to a 6 GB-VRAM GPU, which constrains batch sizes and "
    "fine-tuning depth.",
    "The system operates offline as a research pipeline and is not designed for real-time "
    "or production deployment.",
]
for i, l in enumerate(limitations, 1):
    runs_para([(f"{i}. ", False, False), (l, False, False)], indent_left=0.6)

# ── 3. Methods Used in the Project ────────────────────────────────────────────
heading("3. Methods Used in the Project")

runs_para([("3.1. Gathering Tools Used", True, False)], space_after=6)
para(
    "Because DeepSentinel is a data-driven research system, data gathering relied on "
    "documentary and software-based tools rather than human respondents:",
    indent_left=0.3,
)
gathering = [
    "Library and Literature Research — peer-reviewed studies on deepfake detection, "
    "multimodal emotion recognition, and cross-modal coherence were reviewed to justify "
    "the architecture and to source baseline comparisons.",
    "Documentary Analysis of Benchmark Datasets — established public datasets supplied "
    "the raw material: CREMA-D (source for generated fakes), MELD (real clips and the "
    "Track 4 fake source), CMU-MOSEI (in-the-wild real clips), MUStARD (sarcasm labels), "
    "and FakeAVCeleb (unseen evaluation set).",
    "Software Experimentation and Observation — pretrained models obtained from the "
    "Hugging Face repository (Wav2Vec2, BERT, ViT) and open-source generators (StyleTTS2, "
    "RVC, Wav2Lip, SadTalker, MuseTalk) were used to produce and process data; "
    "experimental runs, training logs, and evaluation metrics were observed and recorded.",
    "Development Toolchain — Python, PyTorch, the Hugging Face Transformers library, "
    "ffmpeg, and InsightFace formed the implementation environment in which the system "
    "was constructed and tested.",
]
for i, g in enumerate(gathering, 1):
    runs_para([(f"{i}. ", False, False), (g, False, False)], indent_left=0.6)

runs_para([("3.2. Engineering Paradigm / Process Model", True, False)], space_after=6)
para(
    "The project adopts the Iterative and Incremental Process Model. This paradigm is "
    "justified by the experimental nature of machine-learning development, in which a "
    "complete solution cannot be specified up front and must instead be refined through "
    "repeated cycles of build, measure, and improve.",
    indent_left=0.3,
)
para(
    "Each iteration delivered an increment of working capability—first the four "
    "generation tracks, then the feature-extraction pipeline, then the detection model, "
    "and finally the evaluation suite. Within the detection model itself, the two-phase "
    "training strategy mirrors the iterative philosophy: Phase 1 establishes a working "
    "classifier on cached features, and Phase 2 incrementally improves it by fine-tuning "
    "the encoders. Empirical feedback at every cycle—loss curves, validation accuracy, "
    "and benchmark AUC—directed the next round of refinement, allowing the team to detect "
    "problems early and to converge on a configuration that generalizes to unseen "
    "manipulation methods. This makes the Iterative and Incremental model a more faithful "
    "fit than a strictly linear Waterfall approach, which assumes fixed, fully understood "
    "requirements from the outset.",
    indent_left=0.3,
)

# ── 4. Project Organizational Structure ───────────────────────────────────────
heading("4. Project Organizational Structure")
para(
    "The team is organized around the system's major modules, with one member designated "
    "as project leader. Roles are summarized below; member names are to be supplied by the "
    "team.",
)
org_tbl = doc.add_table(rows=1, cols=3)
org_tbl.style = "Table Grid"
org_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = org_tbl.rows[0].cells
for c, txt in zip(hdr, ["Name", "Role / Position", "Primary Responsibility"]):
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.paragraphs[0].add_run(txt)
    _set_run(r, bold=True)
    _shade_cell(c, "D9D9D9")
org_rows = [
    ("[LEADER NAME]", "Project Leader", "Overall coordination, integration, and Detection Module"),
    ("[MEMBER NAME]", "Data Generation Lead", "Tracks 1–4 deepfake generation pipeline"),
    ("[MEMBER NAME]", "Preprocessing Lead", "Feature extraction (Wav2Vec2, BERT, ViT) and caching"),
    ("[MEMBER NAME]", "Training Lead", "Two-phase training and multi-task loss"),
    ("[MEMBER NAME]", "Evaluation & Documentation Lead", "Benchmarking, statistical testing, and documentation"),
]
for name, role, resp in org_rows:
    cells = org_tbl.add_row().cells
    for cell, val, center in ((cells[0], name, True), (cells[1], role, False), (cells[2], resp, False)):
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        r = cell.paragraphs[0].add_run(val)
        _set_run(r)

# ── 5. Project Schedule ───────────────────────────────────────────────────────
heading("5. Project Schedule")
para(
    "The project was carried out over the duration of the term. The schedule below lists "
    "the major activities from start to finish; specific dates are to be supplied by the "
    "team.",
)
sched_tbl = doc.add_table(rows=1, cols=4)
sched_tbl.style = "Table Grid"
sched_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
shdr = sched_tbl.rows[0].cells
for c, txt in zip(shdr, ["Phase / Activity", "Start Date", "End Date", "Deliverable"]):
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.paragraphs[0].add_run(txt)
    _set_run(r, bold=True)
    _shade_cell(c, "D9D9D9")
# empty rows for the team to fill
for _ in range(6):
    cells = sched_tbl.add_row().cells
    for cell in cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── 6. Problems Encountered ───────────────────────────────────────────────────
heading("6. Problems Encountered")

# Shaded title bar
title_tbl = doc.add_table(rows=2, cols=1)
title_tbl.style = "Table Grid"
title_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tc0 = title_tbl.rows[0].cells[0]
_shade_cell(tc0, "D9D9D9")
r = tc0.paragraphs[0].add_run("PROBLEMS ENCOUNTERED")
_set_run(r, bold=True, italic=True)
tc1 = title_tbl.rows[1].cells[0]
r = tc1.paragraphs[0].add_run(
    "Identify the problems in analysis, design, development, group composition and other "
    "pertinent concerns."
)
_set_run(r, italic=True)

# 3-column empty table
prob_tbl = doc.add_table(rows=1, cols=3)
prob_tbl.style = "Table Grid"
prob_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
phdr = prob_tbl.rows[0].cells
for c, txt in zip(phdr, ["Problems Encountered", "Solution", "Person Responsible"]):
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.paragraphs[0].add_run(txt)
    _set_run(r, bold=True)
for _ in range(5):
    cells = prob_tbl.add_row().cells
    for cell in cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Footer page numbers (bottom center) ───────────────────────────────────────
_add_page_number_footer()

doc.save(OUT)
print("Saved:", OUT)
print("Paragraphs:", len(doc.paragraphs), "Tables:", len(doc.tables))
